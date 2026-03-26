import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
doc = Document(r'c:\Users\alejandro.carvajal\Documents\Formularios\backend\analisis\Requerimientos_Modulo_Analisis_Psicosocial (1).docx')
for para in doc.paragraphs:
    if para.text.strip():
        style = para.style.name if para.style else 'Normal'
        print(style + ' | ' + para.text)

# Also extract tables
for i, table in enumerate(doc.tables):
    print(f'\n--- TABLE {i+1} ---')
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        print(' | '.join(cells))
